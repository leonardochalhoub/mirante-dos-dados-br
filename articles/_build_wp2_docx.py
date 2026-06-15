# -*- coding: utf-8 -*-
"""Constrói wp2.docx (formato RAP) a partir de wp2-rap.tex.

Pipeline: rasteriza figuras PDF->PNG (PyMuPDF) -> resolve numeração de
seções/tabelas/figuras e \\ref para números estáticos -> converte tcolorbox
-> pandoc latex->docx -> patch styles (Times New Roman 12 / 1,5 / headings)
-> título 14pt + autores 12pt (python-docx).
"""
import os, re, shutil, subprocess, zipfile
from pathlib import Path

ART = Path(__file__).resolve().parent
TEX = ART / "wp2-rap.tex"
OUT = ART / "wp2.docx"
FIGDIR = ART / "figures-pbf"
PANDOC = os.path.expanduser("~/.local/bin/pandoc")
SRC = Path("/tmp/wp2_docx_src.tex")


# ---------- 0. rasteriza PDF -> PNG -----------------------------------
def rasterize():
    import fitz
    for pdf in sorted(FIGDIR.glob("*.pdf")):
        png = pdf.with_suffix(".png")
        doc = fitz.open(pdf)
        pix = doc[0].get_pixmap(matrix=fitz.Matrix(3, 3))  # ~216 dpi
        pix.save(png)
        doc.close()
    print("PNG geradas:", len(list(FIGDIR.glob("*.png"))))


# ---------- util: chaves balanceadas ----------------------------------
def find_balanced(s, i):  # i aponta para '{'
    depth = 0
    while i < len(s):
        if s[i] == "{":
            depth += 1
        elif s[i] == "}":
            depth -= 1
            if depth == 0:
                return i
        i += 1
    return -1


# ---------- 1. numeração estática + cross-refs -------------------------
def resolve_numbers(s):
    # mapa label -> número
    sec = sub = tab = fig = 0
    last_heading = ""
    label_map = {}
    tok = re.compile(
        r"\\section\{|\\subsection\{|\\begin\{table\}|\\begin\{figure\}|\\label\{([^}]*)\}")
    for m in tok.finditer(s):
        g = m.group(0)
        if g == r"\section{":
            sec += 1; sub = 0; last_heading = str(sec)
        elif g == r"\subsection{":
            sub += 1; last_heading = f"{sec}.{sub}"
        elif g == r"\begin{table}":
            tab += 1
        elif g == r"\begin{figure}":
            fig += 1
        else:
            lab = m.group(1)
            if lab.startswith("tab:"):
                label_map[lab] = str(tab)
            elif lab.startswith("fig:"):
                label_map[lab] = str(fig)
            elif lab.startswith("sec:"):
                label_map[lab] = last_heading
    # prefixa legendas (Tabela N — / Figura N —) na ordem dos ambientes
    begins = [(m.start(), "tab" if "table" in m.group(0) else "fig")
              for m in re.finditer(r"\\begin\{table\}|\\begin\{figure\}", s)]
    tn = fn = 0
    numbered = []
    for pos, kind in begins:
        if kind == "tab":
            tn += 1; numbered.append((pos, "Tabela", tn))
        else:
            fn += 1; numbered.append((pos, "Figura", fn))

    out, i = [], 0
    cap = "\\caption{"
    while True:
        j = s.find(cap, i)
        if j < 0:
            out.append(s[i:]); break
        out.append(s[i:j])
        b = j + len(cap) - 1
        e = find_balanced(s, b)
        inner = s[b + 1:e]
        # ambiente do caption = último \begin antes de j
        kind, num = "Figura", fn
        for pos, k, n in numbered:
            if pos < j:
                kind, num = k, n
            else:
                break
        out.append(f"\\caption{{{kind} {num} --- {inner}}}")
        i = e + 1
    s = "".join(out)

    # substitui \ref{...} pelos números; remove \label{...}
    s = re.sub(r"\\ref\{([^}]*)\}", lambda m: label_map.get(m.group(1), "?"), s)
    s = re.sub(r"\\label\{[^}]*\}", "", s)
    return s


# ---------- 1b. achata figuras: legenda(topo) -> imagem -> fonte ------
def flatten_figures(s):
    out, i = [], 0
    beg, end = r"\begin{figure}", r"\end{figure}"
    while True:
        a = s.find(beg, i)
        if a < 0:
            out.append(s[i:]); break
        out.append(s[i:a])
        b = s.find(end, a)
        body = s[a:b]
        img = re.search(r"\\includegraphics(\[[^\]]*\])?\{[^}]*\}", body)
        imgtex = img.group(0) if img else ""
        cap = ""
        ci = body.find(r"\caption{")
        if ci >= 0:
            cs = ci + len(r"\caption")
            ce = find_balanced(body, cs)
            cap = body[cs + 1:ce]
        fon = ""
        fi = body.find(r"{\small Fonte")
        if fi >= 0:
            fe = find_balanced(body, fi)
            fon = re.sub(r"^\\small\s*", "", body[fi + 1:fe])
        blk = "\n\n"
        if cap:
            blk += r"\textit{" + cap + "}" + "\n\n"   # legenda no topo
        blk += imgtex + "\n\n"
        if fon:
            blk += fon + "\n\n"                         # fonte embaixo
        out.append(blk)
        i = b + len(end)
    return "".join(out)


# ---------- 2. tcolorbox -> parágrafo ---------------------------------
def unbox_tcolorbox(s):
    out, i = [], 0
    while True:
        j = s.find(r"\begin{tcolorbox}", i)
        if j < 0:
            out.append(s[i:]); break
        out.append(s[i:j])
        # pula opções [ ... ] (com chaves balanceadas internas)
        k = s.find("[", j)
        end_opt = s.find(r"\end{tcolorbox}", j)
        title = ""
        mt = re.search(r"title=\{(.*?)\}\,?\s*\n", s[j:end_opt], re.S)
        # corpo após o "]" de opções
        depth, p = 0, k
        while p < len(s):
            if s[p] == "[":
                depth += 1
            elif s[p] == "]":
                depth -= 1
                if depth == 0:
                    break
            p += 1
        body = s[p + 1:end_opt].strip()
        body = re.sub(r"^\\small\s*", "", body)
        out.append("\n\n\\textbf{Nota metodológica.} " + body + "\n\n")
        i = end_opt + len(r"\end{tcolorbox}")
    return "".join(out)


# ---------- 3. preprocess geral --------------------------------------
def anonymize(s):
    """Remove identificação dos autores (RAP é double-blind)."""
    s = re.sub(r"(?m)^%.*$", "", s)  # comentários (têm nome do autor)

    def rep(m):
        body = m.group(0)
        if any(k in body for k in ("Chalhoub", "ORCID", "Autor correspondente")):
            return ""
        return body
    s = re.sub(r"\\begin\{center\}.*?\\end\{center\}", rep, s, flags=re.S)
    note = (r"\begin{center}{\footnotesize\itshape [Identificação dos autores "
            r"removida para avaliação por pares duplo-cega.]\par}\end{center}")
    s = s.replace(r"\section*{Resumo}", note + "\n\n" + r"\section*{Resumo}", 1)
    return s


def preprocess(anon=False):
    s = TEX.read_text(encoding="utf-8")
    if anon:
        s = anonymize(s)
    s = s.replace(r"\input{compile-stamp}", "")
    s = s.replace(r" (\COMPILESHA)", "")
    # math que o pandoc não digere (acentos em \mathit, \rs dentro de math)
    s = s.replace(r"\mathit{penetra\c{c}\~ao}", r"\mathit{penetracao}")
    s = s.replace(r"\text{-}", "-")
    s = s.replace(r"$y_{\mathrm{Cad\acute{U}n}} \lessgtr \rs 218$",
                  r"renda do CadÚnico $\gtrless$ R\$~218")
    # pandoc 3.2 descarta as macros \textendash/\textemdash -> usa unicode
    s = s.replace(r"\textemdash{}", "—").replace(r"\textemdash ", "— ")
    s = s.replace(r"\textemdash", "—")
    s = s.replace(r"\textendash{}", "–").replace(r"\textendash ", "–")
    s = s.replace(r"\textendash", "–")
    s = resolve_numbers(s)
    s = flatten_figures(s)
    s = unbox_tcolorbox(s)
    # refitem -> separador de parágrafo
    s = re.sub(r"\\newcommand\{\\refitem\}.*?\n", "", s)
    s = s.replace(r"\refitem", "\n\n")
    # singlespace -> nada
    s = s.replace(r"\begin{singlespace}", "").replace(r"\end{singlespace}", "")
    # figuras: \includegraphics[..]{figXX} -> figures-pbf/figXX.png
    s = re.sub(r"\\includegraphics(\[[^\]]*\])?\{(fig\d+[^}]*)\}",
               r"\\includegraphics\1{figures-pbf/\2.png}", s)
    # remove comandos só de layout que poluem o pandoc
    s = re.sub(r"\\vspace\*?\{[^}]*\}", "", s)
    s = re.sub(r"\\addcontentsline\{[^}]*\}\{[^}]*\}\{[^}]*\}", "", s)
    for c in (r"\clearpage", r"\newpage", r"\par", r"\noindent", r"\centering",
              r"\onehalfspacing", r"\FloatBarrier"):
        s = s.replace(c, "")
    SRC.write_text(s, encoding="utf-8")
    return SRC


# ---------- 4. pandoc -------------------------------------------------
def run_pandoc(src):
    subprocess.run([PANDOC, str(src), "-f", "latex", "-t", "docx",
                    "--resource-path", f"{ART}:{FIGDIR}",
                    "-o", str(OUT)], check=True, cwd=ART)


# ---------- 5. patch styles.xml (TNR12 / 1,5 / headings) -------------
def rpr(bold=True, caps=False):
    b = "<w:b /><w:bCs />" if bold else ""
    c = "<w:caps />" if caps else ""
    return ('<w:rPr><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman" '
            'w:cs="Times New Roman" />' + b + c +
            '<w:color w:val="000000" /><w:sz w:val="24" /><w:szCs w:val="24" /></w:rPr>')


def patch_styles():
    z = zipfile.ZipFile(OUT)
    x = z.read("word/styles.xml").decode("utf-8")
    x = x.replace(
        '<w:rFonts w:asciiTheme="minorHAnsi" w:cstheme="minorBidi" w:eastAsiaTheme="minorHAnsi" w:hAnsiTheme="minorHAnsi" />',
        '<w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:cs="Times New Roman" />')
    x = x.replace(
        '<w:pPrDefault>\n      <w:pPr>\n        <w:spacing w:after="200" />\n      </w:pPr>\n    </w:pPrDefault>',
        '<w:pPrDefault>\n      <w:pPr>\n        <w:spacing w:after="0" w:line="360" w:lineRule="auto" />\n        <w:jc w:val="both" />\n      </w:pPr>\n    </w:pPrDefault>')

    def ph(x, sid, bold, caps):
        pat = re.compile(r'(<w:style w:styleId="' + sid +
                         r'" w:type="paragraph">.*?)<w:rPr>.*?</w:rPr>(\s*</w:style>)', re.S)
        return pat.sub(lambda m: m.group(1) + rpr(bold, caps) + m.group(2), x)

    x = ph(x, "Heading1", True, True)
    x = ph(x, "Heading2", True, False)
    x = ph(x, "Heading3", False, False)
    for n in (4, 5, 6, 7, 8, 9):
        x = ph(x, f"Heading{n}", True, False)

    tmp = Path("/tmp/_wp2_docx_work")
    shutil.rmtree(tmp, ignore_errors=True)
    z.extractall(tmp)
    (tmp / "word" / "styles.xml").write_text(x, encoding="utf-8")
    OUT.unlink()
    zo = zipfile.ZipFile(OUT, "w", zipfile.ZIP_DEFLATED)
    for root, _, files in os.walk(tmp):
        for fn in files:
            full = Path(root) / fn
            zo.write(full, full.relative_to(tmp))
    zo.close()


# ---------- 6. título 14pt + autores 12pt ----------------------------
def finish():
    import docx
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    from docx.shared import Cm
    d = docx.Document(str(OUT))

    # A4 + margens RAP (top/bottom 2,5 cm; left/right 3 cm)
    sec = d.sections[0]
    sec.page_width, sec.page_height = Cm(21), Cm(29.7)
    sec.top_margin = sec.bottom_margin = Cm(2.5)
    sec.left_margin = sec.right_margin = Cm(3)

    def set_run(r, size, italic=None, bold=None):
        r.font.name = "Times New Roman"; r.font.size = Pt(size)
        if italic is not None:
            r.font.italic = italic
        if bold is not None:
            r.font.bold = bold
        rprx = r._element.get_or_add_rPr()
        rf = rprx.find(qn("w:rFonts"))
        if rf is None:
            rf = OxmlElement("w:rFonts"); rprx.append(rf)
        for a in ("w:ascii", "w:hAnsi", "w:cs"):
            rf.set(qn(a), "Times New Roman")

    from docx.shared import Cm, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # figuras: largura plena da mancha (15 cm) preservando proporção
    TEXTW = Cm(15)
    for shp in d.inline_shapes:
        if not shp.width:
            continue
        asp = shp.height / shp.width
        shp.width = TEXTW
        shp.height = Emu(int(int(TEXTW) * asp))
    for p in d.paragraphs:
        if p._p.findall(".//" + qn("w:drawing")):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    in_body = False
    for p in d.paragraphs:
        tl = p.text.lower()
        st = p.style.name
        if "três regimes, um programa" in tl and "documentação" in tl:
            for r in p.runs:
                set_run(r, 14, bold=True)
        elif "three regimes, one program" in tl or "tres regímenes, un programa" in tl:
            for r in p.runs:
                set_run(r, 12, italic=True)
        elif "leonardo chalhoub" in tl and "jefferson" in tl:
            for r in p.runs:
                set_run(r, 12)
        # linha "Fonte:" sob figuras/tabelas -> centralizada, itálica, 10 pt
        if tl.startswith("fonte"):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            for r in p.runs:
                set_run(r, 10, italic=True)
            continue
        # legenda "Figura N — ..." / "Tabela N — ..." -> centralizada, sem recuo
        if re.match(r"(figura|tabela) \d+ [—-]", tl):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            for r in p.runs:
                set_run(r, 10, italic=True)
            continue
        # recuo de 1,25 cm na primeira linha — só no corpo (Introdução..Referências)
        if st.startswith("Heading") and tl.startswith("introdução"):
            in_body = True
        if st.startswith("Heading") and tl.startswith("referências"):
            in_body = False
        if in_body and not st.startswith("Heading") and not st.startswith("Image") \
                and p.text.strip():
            p.paragraph_format.first_line_indent = Cm(1.25)
    d.save(str(OUT))


def build(anon):
    global OUT, SRC
    OUT = ART / ("wp2-anonimo.docx" if anon else "wp2.docx")
    SRC = Path("/tmp/wp2_docx_src%s.tex" % ("_anon" if anon else ""))
    run_pandoc(preprocess(anon))
    patch_styles()
    finish()
    print("built", OUT.name, f"({OUT.stat().st_size//1024} KB)")


if __name__ == "__main__":
    rasterize()
    build(anon=False)   # versão identificada (registro/versão final)
    build(anon=True)    # versão cega (upload no ScholarOne)
    # .tex anônimo compilável (LaTeX limpo, sem transformações de docx)
    texout = ART / "wp2-anonimo.tex"
    texout.write_text(anonymize(TEX.read_text(encoding="utf-8")), encoding="utf-8")
    print("built", texout.name)
