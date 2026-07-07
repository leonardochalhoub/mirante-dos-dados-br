#!/usr/bin/env python3
"""Monta o manuscrito "O cálculo ausente" DENTRO do template oficial do Bolema.

Motivo: a submissão BOLEMA-2026-0250 foi reprovada na triagem porque o .docx
fora gerado por conversão LaTeX->Word (estilos do pandoc), e não digitado no
template exigido (TEMPLATE_BOLEMA_PT-2025.docx). Este script parte do template
oficial, limpa os placeholders, e despeja o conteúdo usando os ESTILOS do
próprio template. Fonte de conteúdo: calculo-bolema-submissao.docx (texto +
figuras já resolvidos), com correções: PT+EN apenas (sem ES), matemática do
Quadro 1 e dos exemplos de exame recuperada do .tex, e fontes das figuras.

Reproduzível: python articles/scripts/build_docx_bolema_template.py
"""
import io
from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH as AL
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT as VA
from docx.oxml.ns import qn

TEMPLATE = "TEMPLATE_BOLEMA_PT-2025.docx"
SRC = "articles/calculo-bolema-submissao.docx"
OUT = "articles/calculo-bolema-TEMPLATE.docx"
MAXW = Cm(10.0)  # figuras compactas (mantêm legibilidade; poupam páginas)

T = Document(TEMPLATE)
S = Document(SRC)
STYLES = {s.name for s in T.styles}


def tnr(run, size=None, bold=None, italic=None):
    """Força Times New Roman (ascii/hAnsi/cs/eastAsia) e atributos."""
    run.font.name = "Times New Roman"
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = rpr.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, rf)
    for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rf.set(qn(a), "Times New Roman")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    return run


def par(style, align=None):
    p = T.add_paragraph(style=style if style in STYLES else None)
    if align is not None:
        p.alignment = align
    return p


def line(text, style="Texto Comum", align=None, size=None, bold=None, italic=None):
    p = par(style, align)
    tnr(p.add_run(text), size, bold, italic)
    return p


def copy_runs(src_p, style, size=None, align=None, force_bold=None):
    """Copia runs preservando itálico/negrito (refs, termos estrangeiros)."""
    p = par(style, align)
    runs = src_p.runs
    if not runs:
        return p
    for sr in runs:
        if not sr.text:
            continue
        r = p.add_run(sr.text)
        tnr(r, size, force_bold if force_bold is not None else sr.bold, sr.italic)
    return p


# ---------------------------------------------------------------- limpar corpo
body = T.element.body
sectPr = body.find(qn("w:sectPr"))  # mantido no corpo p/ add_table calcular largura
for ch in list(body):
    if ch is not sectPr:
        body.remove(ch)

# ---------------------------------------------------------------- folha de rosto
SP = S.paragraphs

def space(p, before=0, after=0):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    return p

space(line(SP[0].text, "Normal", AL.CENTER, size=16, bold=True), after=6)   # título PT
space(line(SP[4].text, "Normal", AL.CENTER, size=14, bold=True), after=12)  # título EN
# (autores omitidos: versão anônima para avaliação duplo-cega)

# Bloco de resumo conforme o template: cabeçalho CENTRALIZADO bold; texto e
# palavras-chave JUSTIFICADOS em TNR 10 (o estilo 'local' é Arial 14 por baixo —
# por isso fonte e tamanho são forçados explicitamente).
space(line("Resumo", "Normal", AL.CENTER, bold=True), before=6)
line(SP[2].text, "local", AL.JUSTIFY, size=10)                              # resumo PT (TNR 10)
space(line(SP[3].text, "local", AL.JUSTIFY, bold=True, size=10), before=6, after=6)   # palavras-chave PT
space(line("Abstract", "Normal", AL.CENTER, bold=True), before=6)
line(SP[6].text, "local", AL.JUSTIFY, size=10)                             # abstract EN (TNR 10)
space(line(SP[7].text, "local", AL.JUSTIFY, bold=True, size=10), before=6, after=12)  # keywords EN

# ---------------------------------------------------------------- corpo
FIG_SRC = {
    1: "Fonte: Elaboração própria, a partir de PNUD (2024) e dos currículos oficiais analisados.",
    2: "Fonte: Elaboração própria.",
    3: "Fonte: Elaboração própria, a partir de OECD (2023).",
}
fig_n = 0

# Quadro 1 com matemática corrigida (recuperada do .tex)
QUADRO1 = [
    ["Modo", "Situação (derivada / integral)", "Objeto matemático"],
    ["Enativo\n(concreto)",
     "Móvel: medir posição e tempo (velocidade); e recuperar a distância percorrida a partir da velocidade",
     "Taxa média Δs/Δt; distância como área acumulada"],
    ["Icônico\n(gráfico)",
     "Gráfico posição–tempo (secantes → tangente); e área sob o gráfico velocidade–tempo por retângulos",
     "Tangente como limite de secantes; área como limite de somas de retângulos"],
    ["Simbólico\n(formal)",
     "Razão incremental → limite; e somas de Riemann → integral definida",
     "f′(x)=lim(h→0)[f(x+h)−f(x)]/h e ∫ₐᵇ f; o TFC liga as duas"],
]


def build_quadro():
    tb = T.add_table(rows=len(QUADRO1), cols=3)
    tb.style = "Table Grid" if "Table Grid" in STYLES else None
    tb.alignment = 1  # center
    for i, row in enumerate(QUADRO1):
        for j, txt in enumerate(row):
            cell = tb.cell(i, j)
            cell.vertical_alignment = VA.CENTER          # centraliza vertical
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = AL.CENTER                       # centraliza horizontal
            p.paragraph_format.first_line_indent = Cm(0)  # sem recuo herdado
            p.paragraph_format.left_indent = Cm(0)
            p.paragraph_format.right_indent = Cm(0)
            parts = txt.split("\n")
            for k, seg in enumerate(parts):
                if k:
                    p.add_run().add_break()
                tnr(p.add_run(seg), size=10, bold=(i == 0))


def emit_figure(child):
    global fig_n
    fig_n += 1
    blip = child.find(".//" + qn("a:blip"))
    rid = blip.get(qn("r:embed"))
    blob = S.part.related_parts[rid].blob
    ext = child.find(".//" + qn("wp:extent"))
    cx = int(ext.get("cx")) if ext is not None else int(MAXW)
    width = Emu(min(cx, int(MAXW)))
    p = par("Texto ABNT", AL.CENTER)
    p.paragraph_format.first_line_indent = Cm(0)  # sem recuo: centraliza de fato
    p.paragraph_format.left_indent = Cm(0)
    p.add_run().add_picture(io.BytesIO(blob), width=width)


# iteração ordenada sobre os filhos do corpo de S
capture = False
pending_lead = None
skip_table = False

for child in S.element.body.iterchildren():
    tag = child.tag.split("}")[-1]
    if tag == "tbl":
        if skip_table:
            skip_table = False
            build_quadro()
        continue
    if tag != "p":
        continue
    from docx.text.paragraph import Paragraph
    p = Paragraph(child, S)
    style = p.style.name
    text = p.text.strip()

    if not capture:
        if style == "Heading 1" and text.replace("\t", " ").strip().startswith("1"):
            capture = True
        else:
            continue
    # parar no bloco final (tratado à parte)
    if style == "Heading 1" and text.startswith("Agradecimentos"):
        break

    is_img = bool(child.findall(".//" + qn("w:drawing")))

    if is_img:
        emit_figure(child)
        continue
    if not text:
        continue

    # legenda de figura -> Texto ABNT bold 10 centralizado + Fonte
    if style == "Image Caption" or text.startswith("Figura "):
        cap = line(text, "Texto ABNT", AL.CENTER, size=10, bold=True)
        cap.paragraph_format.first_line_indent = Cm(0)
        cap.paragraph_format.left_indent = Cm(0)
        if fig_n in FIG_SRC:
            src = line(FIG_SRC[fig_n], "Texto ABNT", AL.CENTER, size=10)
            src.paragraph_format.first_line_indent = Cm(0)
            src.paragraph_format.left_indent = Cm(0)
        continue
    # legenda de quadro -> Caption + tabela nativa
    if style == "Table Caption" or text.startswith("Quadro "):
        line(text, "Caption", size=10, bold=True)
        skip_table = True
        continue
    # títulos de seção/subseção -> Normal bold 12
    if style in ("Heading 1", "Heading 2", "Heading 3"):
        line(text.replace("\t", " "), "Normal", bold=True, size=12)
        continue
    # rótulos regionais (Heading 4: "4.1.0.1 Ásia.") -> lead-in em negrito
    if style == "Heading 4":
        lab = text.split("\t")[-1].strip()
        pending_lead = lab if lab.endswith(".") else lab + "."
        continue
    # citação longa (>3 linhas) -> recuo ABNT, TNR 10
    if style == "Block Text":
        q = par("Normal", AL.JUSTIFY)
        q.paragraph_format.left_indent = Cm(4)
        q.paragraph_format.line_spacing = 1.0
        for sr in (p.runs or [p]):
            tnr(q.add_run(sr.text), size=10, italic=getattr(sr, "italic", None))
        continue
    # Fonte do quadro
    if text.startswith("Fonte:"):
        line(text, "Texto Comum", size=10)
        continue

    # corpo comum (com possível lead-in regional em negrito)
    tp = par("Texto Comum", AL.JUSTIFY)
    if pending_lead:
        tnr(tp.add_run(pending_lead + " "), bold=True)
        pending_lead = None
    for sr in p.runs:
        if sr.text:
            tnr(tp.add_run(sr.text), italic=sr.italic, bold=sr.bold)
    if not p.runs:
        tnr(tp.add_run(text))

# ------------------ correção de \ref e matemática dropados pelo pandoc --------
# O docx-fonte (pandoc) descartou TODO \ref e toda matemática inline. Recuperados
# do .tex. Números de referência: fig:hdi=1, fig:heatmap=2, fig:pisa=3,
# quad:cpa=1, sec:metodo=3, sec:sintese=4.3, sec:discussao=4.8.

def fix_runs(p, segments):
    """Substitui os runs de p pelos segmentos (texto, itálico), preservando estilo."""
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    for txt, ital in segments:
        tnr(p.add_run(txt), italic=ital or None)

# (a) rewrites ricos (preservam ênfases) -------------------------------------
SPEARMAN = [
    ("Uma verificação estatística acompanha esse panorama. Codificando o status "
     "curricular de forma ordinal (obrigatório = 2; eletivo/opção de trilha avançada "
     "= 1; ausente = 0 — codificação documentada no script ", False),
    ("compute_pisa_correlations.py", True),
    (") e correlacionando-o com o desempenho no PISA 2022 dos quinze sistemas com dado "
     "disponível, o coeficiente de correlação de postos de Spearman é nulo (ρ = −0,08; "
     "IC 95% por ", False), ("bootstrap", True),
    (" [−0,67, 0,51]; teste de permutação unilateral p ≈ 0,61, hipótese de gradiente "
     "positivo). A não-significância é estável sob codificações ordinais alternativas do "
     "status curricular. Isto é: entre os países que ", False), ("têm", True),
    (" cálculo, o grau de obrigatoriedade não prediz o desempenho — resultado coerente "
     "com a ressalva de que o PISA é um ", False), ("proxy", True),
    (" invertido (Seção 4.8). O que é robusto não é um gradiente dose-resposta, e sim a "
     "posição do Brasil como ", False), ("outlier", True),
    (": seu escore (379) está a cerca de 3,5 desvios-padrão abaixo da média dos sistemas "
     "com cálculo (492). O contraste ausente-", False), ("versus", True),
    ("-presente também é marcante (correlação ponto-bisserial r = 0,68; p < 0,01), embora, "
     "com o Brasil como único caso “ausente”, esse coeficiente seja uma reparametrização "
     "do mesmo escore-z, não uma evidência independente. O código que reproduz estas "
     "estatísticas é público (ver Declarações). Os dados sustentam, portanto, a ", False),
    ("singularidade curricular", True),
    (" do Brasil — não uma relação causal entre cálculo no ensino médio e desempenho aos "
     "15 anos.", False),
]
EXAMS = [
    ("Gaokao (China):", False, True),
    (" seja f(x)=x³−3ax+2; discuta a monotonicidade de f, determine a para que o mínimo "
     "seja 0 e o número de raízes de f(x)=1 — ", False, False),
    ("exige derivar e analisar variação", True, False),
    (" (análogos formais constam do Abitur alemão e do AP Calculus BC). ", False, False),
    ("ENEM (Brasil):", False, True),
    (" o número de bactérias dobra a cada hora; a partir de 100, quantas haverá após 5 "
     "horas? — ", False, False),
    ("crescimento exponencial em forma discreta, sem interface com a derivada nem com a "
     "taxa instantânea", True, False),
    (".", False, False),
]
for p in T.paragraphs:
    if "Uma verificação estatística acompanha" in p.text:
        fix_runs(p, SPEARMAN)
    elif "monotonicidade" in p.text and "discuta a" in p.text:
        for r in list(p.runs):
            r._element.getparent().remove(r._element)
        for txt, ital, bold in EXAMS:
            r = p.add_run(txt); tnr(r, italic=ital or None, bold=bold or None)

# (b) correções de token (parágrafos sem ênfase relevante) --------------------
TOKEN = [
    ("retorna na discussão (Seção ).", "retorna na discussão (Seção 4.8)."),
    ("verificável de forma automatizada (Seção ).", "verificável de forma automatizada (Seção 3)."),
    ("O Quadro  organiza", "O Quadro 1 organiza"),
    ("demonstradas (Seção ); e (iv)", "demonstradas (Seção 4.8); e (iv)"),
    ("consta da Figura  e do repositório", "consta da Figura 2 e do repositório"),
    ("A Figura  torna visual", "A Figura 1 torna visual"),
    ("A Figura  sintetiza", "A Figura 2 sintetiza"),
    ("célula a célula (país  tópico)", "célula a célula (país × tópico)"),
    ("canal causal discutido na Seção .", "canal causal discutido na Seção 4.8."),
    ("estimar seu efeito (Seção ).", "estimar seu efeito (Seção 4.8)."),
    ("de 143 pontos em 2003 a 93 em 2022 (OECD, 2023; Figura )",
     "de −143 pontos em 2003 a −93 em 2022 (OECD, 2023; Figura 3)"),
    ("no PISA (; Seção )", "no PISA (ρ = −0,08; Seção 4.3)"),
]
for p in T.paragraphs:
    if "Uma verificação estatística acompanha" in p.text or "monotonicidade" in p.text:
        continue
    probe = p.text.replace(chr(160), " ")  # tex usa ~ (nbsp/U+00A0) antes de \ref
    new = probe
    for broken, fixed in TOKEN:
        new = new.replace(broken, fixed)
    if new != probe:
        fix_runs(p, [(new, False)])

# ---------------------------------------------------------------- seções finais (template)
line("Agradecimentos", "Normal", bold=True)
line("[Suprimidos para avaliação por pares duplo-cega; serão inseridos na versão "
     "final, conforme a política da revista.]", "Texto Comum", AL.JUSTIFY)

line("Contribuições de autoria", "Normal", bold=True)
line("Todos os autores contribuíram substancialmente para a concepção e o desenho "
     "do estudo; para a curadoria e a verificação das fontes curriculares e de "
     "avaliação oficiais; e para a redação e a revisão crítica do manuscrito. Todos "
     "aprovaram a versão submetida e por ela se responsabilizam publicamente.",
     "Texto Comum", AL.JUSTIFY)

line("Disponibilidade de dados", "Normal", bold=True)
line("Os dados analisados provêm integralmente de documentos curriculares e de "
     "avaliação públicos e oficiais, referenciados ao longo do texto. O código e o "
     "manifesto de fontes (com URLs e somas de verificação SHA-256) que reproduzem "
     "o achado central estão depositados em repositório público, cujo endereço será "
     "informado na versão final; durante a avaliação duplo-cega, um espelho "
     "anonimizado pode ser disponibilizado a pedido dos editores.",
     "Texto Comum", AL.JUSTIFY)

# ---------------------------------------------------------------- referências
line("Referências", "Normal", bold=True)
in_refs = False
for p in S.paragraphs:
    if p.style.name == "Heading 1" and p.text.strip().startswith("Referências"):
        in_refs = True
        continue
    if in_refs and p.text.strip():
        rp = copy_runs(p, "Normal", size=11, align=AL.LEFT)
        rp.paragraph_format.line_spacing = 1.0
        rp.paragraph_format.space_after = Pt(2)

# -------- normalizar recuo de 1ª linha: só o CORPO (Texto Comum) tem recuo ----
# O estilo Normal do template embute recuo de 1,25 cm; em títulos/cabeçalhos
# centralizados e nas referências isso entorta o layout (1ª linha deslocada).
for p in T.paragraphs:
    if p.style.name != "Texto Comum":
        p.paragraph_format.first_line_indent = Cm(0)

# ---------------------------------------------------------------- fechar
if sectPr is not None:
    body.remove(sectPr)
    body.append(sectPr)  # sectPr deve ser o último filho do corpo

# limpar propriedades (anonimização) + salvar
cp = T.core_properties
cp.author = ""
cp.last_modified_by = ""
cp.title = "O cálculo ausente"
cp.comments = ""
T.save(OUT)
print("OK ->", OUT, "| figuras:", fig_n)
