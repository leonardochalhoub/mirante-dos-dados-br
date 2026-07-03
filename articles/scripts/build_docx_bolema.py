"""Build the Bolema .docx from calculo-bolema-submissao.tex — template-conformant.

Pipeline: preprocess LaTeX -> pandoc (--number-sections) -> patch styles
(kill theme fonts -> Times New Roman) -> post-process (python-docx) para
GARANTIR, de forma determinística, a conformidade com TEMPLATE_BOLEMA_PT-2025.docx:
  A4 + margens 3/2/3/2 cm; TNR em todo run; corpo 12pt/1,5/recuo 1,25cm;
  títulos centralizados (16/14pt); palavras-chave reinjetadas (pandoc as perde);
  rótulos "Quadro N -"/"Figura N -" nas legendas; propriedades anonimizadas.

Requires: pandoc + python-docx (dev-env). Run:
    conda activate dev-env
    cd articles && python scripts/build_docx_bolema.py
"""
from __future__ import annotations
import os, re, shutil, subprocess, zipfile
from pathlib import Path

ART = Path(__file__).resolve().parents[1]
TEX = ART / "calculo-bolema-submissao.tex"
OUT = ART / "calculo-bolema-submissao.docx"
PANDOC = shutil.which("pandoc") or os.path.expanduser("~/.local/bin/pandoc")

# Palavras-chave (pandoc perde a linha \noindent{\small...}); reinjetadas no fim.
KEYWORDS = {
    "Palavras-chave:": "Cálculo Diferencial e Integral. Currículo de Matemática. "
                       "Ensino Médio. Educação Matemática Comparada. BNCC.",
    "Keywords:": "Differential and Integral Calculus. Mathematics Curriculum. "
                 "Secondary Education. Comparative Mathematics Education. BNCC.",
    "Palabras clave:": "Cálculo Diferencial e Integral. Currículo de Matemática. "
                       "Enseñanza Secundaria. Educación Matemática Comparada. BNCC.",
}


# ---------- 1. preprocess .tex for pandoc -------------------------------
def find_balanced(s, i):
    depth = 0
    while i < len(s):
        if s[i] == '{':
            depth += 1
        elif s[i] == '}':
            depth -= 1
            if depth == 0:
                return i
        i += 1
    return -1


def replace_cmd(s, cmd, fmt):
    out, i, tok = [], 0, "\\" + cmd + "{"
    while True:
        j = s.find(tok, i)
        if j < 0:
            out.append(s[i:]); break
        out.append(s[i:j])
        b = j + len(tok) - 1
        e = find_balanced(s, b)
        out.append(fmt(s[b + 1:e]))
        i = e + 1
    return "".join(out)


def unbox(s):
    out, i, tok = [], 0, "\\fbox{\\parbox{"
    while True:
        j = s.find(tok, i)
        if j < 0:
            out.append(s[i:]); break
        out.append(s[i:j])
        w = j + len("\\fbox{\\parbox")
        we = find_balanced(s, w)
        c = s.find("{", we)
        ce = find_balanced(s, c)
        inner = re.sub(r"^\\small\s*", "", s[c + 1:ce].strip())
        after = ce + 1
        if after < len(s) and s[after] == '}':
            after += 1
        out.append("\n\n" + inner + "\n\n")
        i = after
    return "".join(out)


def preprocess():
    s = TEX.read_text(encoding="utf-8")
    s = re.sub(r"\\newcommand\{\\fonte\}\[1\]\{[^\n]*\}\n", "", s)
    s = re.sub(r"\\newcommand\{\\refitem\}\[1\]\{[^\n]*\}\n", "", s)
    s = replace_cmd(s, "refitem", lambda x: "\n\n" + x.strip() + "\n\n")
    s = replace_cmd(s, "fonte", lambda x: "\n\n\\textit{Fonte: " + x.strip() + "}\n\n")
    s = unbox(s)
    s = re.sub(r"\\includegraphics(\[[^\]]*\])?\{(fig\d+[^}]*)\.pdf\}",
               r"\\includegraphics\1{figures-calculo/\2.png}", s)
    p = Path("/tmp/calculo_bolema_docx_src.tex")
    p.write_text(s, encoding="utf-8")
    return p


# ---------- 2. pandoc -> docx -------------------------------------------
def run_pandoc(src):
    subprocess.run([PANDOC, str(src), "-f", "latex", "-t", "docx",
                    "--number-sections",
                    "--resource-path", f"{ART}:{ART}/figures-calculo",
                    "-o", str(OUT)], check=True, cwd=ART)


# ---------- 3. patch styles.xml: kill theme fonts -> Times New Roman -----
def patch_styles():
    tmp = Path("/tmp/_docx_bolema_work")
    shutil.rmtree(tmp, ignore_errors=True)
    zipfile.ZipFile(OUT).extractall(tmp)
    sp = tmp / "word" / "styles.xml"
    x = sp.read_text(encoding="utf-8")
    # substitui QUALQUER referência a fonte de tema pela Times New Roman
    x = re.sub(r'<w:rFonts[^/]*?(?:asciiTheme|hAnsiTheme|cstheme|eastAsiaTheme)[^/]*?/>',
               '<w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:cs="Times New Roman" />', x)
    # garante docDefaults em TNR + espaçamento 1,5 justificado + idioma pt-BR
    x = re.sub(r'<w:pPrDefault>.*?</w:pPrDefault>',
               '<w:pPrDefault><w:pPr><w:spacing w:after="0" w:line="360" w:lineRule="auto" />'
               '<w:jc w:val="both" /></w:pPr></w:pPrDefault>', x, flags=re.S)
    if '<w:lang' not in x.split('</w:rPrDefault>')[0]:
        x = x.replace('</w:rPrDefault>', '<w:lang w:val="pt-BR" /></w:rPrDefault>', 1)
    # Heading 2 (subseções): não-negrito + versalete (template Bolema)
    def fix_h2(m):
        block = m.group(0)
        block = re.sub(r'<w:b\b[^/]*/>', '', block)
        block = re.sub(r'<w:bCs\b[^/]*/>', '', block)
        if '<w:smallCaps' not in block:
            block = block.replace('</w:rPr>', '<w:smallCaps /></w:rPr>', 1)
        return block
    x = re.sub(r'<w:style w:styleId="Heading2"[^>]*>.*?</w:style>', fix_h2, x, flags=re.S)
    # Heading 1 (seções): CAIXA ALTA (template Bolema)
    def fix_h1(m):
        block = m.group(0)
        if '<w:caps' not in block and '<w:smallCaps' not in block:
            block = block.replace('</w:rPr>', '<w:caps /></w:rPr>', 1)
        return block
    x = re.sub(r'<w:style w:styleId="Heading1"[^>]*>.*?</w:style>', fix_h1, x, flags=re.S)
    sp.write_text(x, encoding="utf-8")
    _anonymize_props(tmp)
    OUT.unlink()
    zo = zipfile.ZipFile(OUT, "w", zipfile.ZIP_DEFLATED)
    for root, _, files in os.walk(tmp):
        for fn in files:
            full = Path(root) / fn
            zo.write(full, full.relative_to(tmp))
    zo.close()


def _anonymize_props(tmp: Path):
    core = tmp / "docProps" / "core.xml"
    if core.exists():
        c = core.read_text(encoding="utf-8")
        for tag in ("dc:creator", "cp:lastModifiedBy", "dc:title"):
            c = re.sub(f"<{tag}>.*?</{tag}>", f"<{tag}></{tag}>", c, flags=re.S)
        core.write_text(c, encoding="utf-8")
    app = tmp / "docProps" / "app.xml"
    if app.exists():
        a = app.read_text(encoding="utf-8")
        for tag in ("Company", "Manager"):
            a = re.sub(f"<{tag}>.*?</{tag}>", f"<{tag}></{tag}>", a, flags=re.S)
        app.write_text(a, encoding="utf-8")


# ---------- 4. post-process (python-docx): enforce template specs --------
def finish():
    import docx
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    d = docx.Document(str(OUT))

    def force_tnr(run):
        run.font.name = "Times New Roman"
        rpr = run._element.get_or_add_rPr()
        rf = rpr.find(qn('w:rFonts'))
        if rf is None:
            rf = OxmlElement('w:rFonts'); rpr.append(rf)
        for a in ('w:ascii', 'w:hAnsi', 'w:cs'):
            rf.set(qn(a), "Times New Roman")

    # (a) A4 + margens 3/2/3/2 cm
    for sec in d.sections:
        sec.orientation = WD_ORIENT.PORTRAIT
        sec.page_width = Cm(21.0); sec.page_height = Cm(29.7)
        sec.top_margin = Cm(3); sec.bottom_margin = Cm(2)
        sec.left_margin = Cm(3); sec.right_margin = Cm(2)

    # localizar títulos e resumos por conteúdo
    def is_title_pt(t): return "cálculo ausente" in t and "duzentos anos" in t
    def is_title_sec(t): return ("absent calculus" in t) or ("doscientos" in t)

    tbl_n = fig_n = 0
    abstract_paras = []  # parágrafos "Resumo/Abstract/Resumen" para injetar keywords após bloco

    for p in d.paragraphs:
        t = p.text.strip()
        tl = t.lower()
        style = (p.style.name or "").lower()

        # fonte TNR em todo run
        for r in p.runs:
            force_tnr(r)

        # (d) títulos centralizados + tamanho
        if is_title_pt(tl):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.font.size = Pt(16); r.font.bold = True
            continue
        if is_title_sec(tl):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.font.size = Pt(14); r.font.bold = True
            continue

        # rótulos de resumo centralizados
        if t in ("Resumo", "Abstract", "Resumen"):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.font.bold = True
            abstract_paras.append(p)
            continue

        # (f) legendas: prefixar "Quadro N -" (Table Caption) / "Figura N -"
        # (Image Caption). "Captioned Figure" é o container da imagem: ignorar.
        if style == "table caption":
            tbl_n += 1; prefix = f"Quadro {tbl_n} – "
        elif style == "image caption":
            fig_n += 1; prefix = f"Figura {fig_n} – "
        else:
            if "caption" in style:
                continue  # Captioned Figure e afins: não rotular
            prefix = None
        if prefix and p.runs and not p.text.startswith(("Figura", "Quadro", "Tabela")):
            run = p.runs[0]
            run.text = prefix + run.text
            run.font.bold = True
            force_tnr(run)
            continue
        if prefix:
            continue

    # (c) recuo de 1a linha 1,25cm SÓ no corpo — nunca em títulos/rótulos
    # centralizados nem em cabeçalhos/legendas/citações.
    for p in d.paragraphs:
        style = (p.style.name or "").lower()
        if any(k in style for k in ("heading", "title", "caption", "quote", "toc")):
            continue
        if p.alignment == WD_ALIGN_PARAGRAPH.CENTER:   # títulos e rótulos de resumo
            continue
        if not p.text.strip():
            continue
        p.paragraph_format.first_line_indent = Cm(1.25)

    # cabeçalhos de seção: TNR 12, PRETO, negrito (H2 = versalete não-negrito).
    # Corrige o azul-petróleo (#0F4761) e o tamanho 10pt herdados do tema.
    from docx.shared import RGBColor
    for p in d.paragraphs:
        sn = (p.style.name or "")
        if sn.startswith("Heading"):
            nonbold = sn.strip().endswith("2")   # Heading 2 = versalete, não-negrito
            for r in p.runs:
                force_tnr(r); r.font.size = Pt(12)
                r.font.color.rgb = RGBColor(0, 0, 0)
                r.font.bold = not nonbold

    # células de tabela: TNR 10 (template)
    for tbl in d.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        force_tnr(r); r.font.size = Pt(10)

    # texto do resumo/abstract/resumen: TNR 10, sem recuo. É APENAS o
    # parágrafo imediatamente após cada rótulo (o resumo é um parágrafo só) ---
    # nunca varre o corpo do artigo (bug anterior zerava o recuo de tudo).
    labels = {"Resumo", "Abstract", "Resumen"}
    paras = d.paragraphs
    for i, p in enumerate(paras):
        if p.text.strip() in labels:
            j = i + 1
            while j < len(paras) and not paras[j].text.strip():
                j += 1
            if j < len(paras):
                for r in paras[j].runs:
                    force_tnr(r); r.font.size = Pt(10)
                paras[j].paragraph_format.first_line_indent = Cm(0)

    d.save(str(OUT))

    # (e) injetar palavras-chave se ausentes (pandoc as perde)
    _inject_keywords(str(OUT))


def _inject_keywords(path):
    import docx
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    d = docx.Document(path)
    full = "\n".join(p.text for p in d.paragraphs)
    if "Palavras-chave" in full:
        return  # já presentes
    # inserir cada bloco de keywords após o parágrafo de resumo correspondente
    lang_after = {"Resumo": "Palavras-chave:", "Abstract": "Keywords:", "Resumen": "Palabras clave:"}
    # localizar o último parágrafo de cada resumo (antes do próximo título/resumo)
    paras = d.paragraphs
    def add_kw_after(anchor_para, label):
        new_p = anchor_para.insert_paragraph_before() if False else None
        # inserir DEPOIS do anchor: cria novo elemento p e insere após
        p_el = OxmlElement('w:p')
        anchor_para._p.addnext(p_el)
        from docx.text.paragraph import Paragraph
        np = Paragraph(p_el, anchor_para._parent)
        r1 = np.add_run(label + " "); r1.bold = True
        r2 = np.add_run(KEYWORDS[label])
        for r in (r1, r2):
            r.font.name = "Times New Roman"; r.font.size = Pt(10)
            rpr = r._element.get_or_add_rPr()
            rf = rpr.find(qn('w:rFonts'))
            if rf is None:
                rf = OxmlElement('w:rFonts'); rpr.append(rf)
            for a in ('w:ascii', 'w:hAnsi', 'w:cs'):
                rf.set(qn(a), "Times New Roman")
        return np
    # estratégia: achar o parágrafo cujo texto começa o resumo e o fim do bloco.
    # Aqui, inserimos as keywords logo após o parágrafo de TEXTO do resumo (o mais
    # longo após o rótulo). Percorremos e guardamos o último parágrafo não-vazio
    # antes de um rótulo de título secundário.
    resumo_labels = ["Resumo", "Abstract", "Resumen"]
    idxs = [i for i,p in enumerate(paras) if p.text.strip() in resumo_labels]
    for k,i in enumerate(idxs):
        # o texto do resumo é o próximo parágrafo não vazio
        j = i+1
        while j < len(paras) and not paras[j].text.strip():
            j += 1
        if j < len(paras):
            label = lang_after[paras[i].text.strip()]
            add_kw_after(paras[j], label)
    d.save(path)


if __name__ == "__main__":
    src = preprocess()
    run_pandoc(src)
    patch_styles()
    finish()
    print("built", OUT)
