#!/usr/bin/env python3
"""Preenche a Cover Letter no TEMPLATE OFICIAL do Bolema (bolema-template-CL.docx).

A carta do Bolema NÃO é uma carta livre: é um questionário obrigatório em três
seções (A: Contextualização 400-500 palavras; B: Rigor Metodológico 400-500;
C: Contribuição 500-600; total <= 1.500). Foi por não seguir essa estrutura que
a 1ª cover letter foi reprovada. Este script parte do template oficial, mantém o
bloco em Português (idioma primário do manuscrito), insere as respostas sob cada
questão e remove os blocos EN/ES não utilizados.

Saída: articles/calculo-bolema-carta-TEMPLATE.docx
"""
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH as AL

TEMPLATE = "TEMPLATE_BOLEMA_CL.docx"
OUT = "articles/calculo-bolema-carta-TEMPLATE.docx"

d = Document(TEMPLATE)

# Respostas (chave = texto exato da última questão do item; valor = resposta) --
A1 = (
    "O manuscrito insere-se em duas subáreas articuladas da Educação Matemática: "
    "(i) currículo e políticas curriculares de Matemática e (ii) Educação Matemática "
    "comparada, com interface direta com a didática do Cálculo Diferencial e Integral. "
    "Os marcos teóricos e metodológicos consolidados que sustentam a análise são: o "
    "princípio do currículo em espiral de Bruner e a sequência Concrete-Pictorial-Abstract "
    "dele derivada; a zona de desenvolvimento proximal de Vygotsky; a transposição "
    "didática de Chevallard, que fornece a categoria central para ler a ausência brasileira "
    "como uma não-transposição do saber de referência ao saber a ensinar; a engenharia "
    "didática de Artigue como método de desenho e análise a priori de sequências; e a "
    "literatura específica sobre a aprendizagem do Cálculo — a distinção entre imagem e "
    "definição conceitual (Tall e Vinner), a reificação (Sfard) e a teoria APOS —, além do "
    "diagnóstico epistemológico de Rezende sobre as dificuldades do ensino de Cálculo no "
    "Brasil. No próprio Bolema, a tradição de estudos curriculares comparados (Pires, Godoy, "
    "Silva e Santos, 2014; Gonçalves, Dias e Peralta, 2018) e a discussão sobre a introdução "
    "das ideias do Cálculo na educação básica (Araújo e Avelar, 2022) delimitam o estado da "
    "arte. Referências fundamentais: Bruner (1960); Chevallard (1991); Artigue (1990); Tall "
    "e Vinner (1981); Rezende (2003); Pires, Godoy, Silva e Santos (2014); Gonçalves, Dias e "
    "Peralta (2018); Araújo e Avelar (2022)."
)
A2 = (
    "Três lacunas justificam a pesquisa. Teórico-empírica: a literatura brasileira reconhece, "
    "há décadas, a ausência do Cálculo no ensino médio (Ávila; Rezende; SBEM), mas sem um "
    "mapeamento internacional preciso e país a país da presença ou ausência do tópico no "
    "currículo prescrito — a afirmação de que “o Brasil é exceção” circula como intuição, "
    "não como achado sistematizado. Comparada: os estudos curriculares já publicados no "
    "periódico tratam sobretudo de estrutura, organização e competências curriculares, não "
    "do recorte específico da presença do Cálculo antes do ensino superior. Metodológica: "
    "afirmações sobre o panorama internacional raramente são acompanhadas de verificação "
    "auditável sobre os documentos oficiais, o que fragiliza sua replicabilidade. Essas "
    "lacunas justificam uma revisão documental comparada, com definição operacional explícita "
    "de “Cálculo no currículo” e verificação reproduzível, capaz de sustentar empiricamente — "
    "e de delimitar com honestidade — o alcance da singularidade brasileira. A lacuna tem, "
    "ainda, consequência prática imediata: sem esse mapeamento, o debate curricular brasileiro "
    "(BNCC e Novo Ensino Médio) discute a reintrodução do Cálculo sem uma referência comparada "
    "firme sobre o que fazem os demais sistemas nacionais."
)
B3 = (
    "A abordagem é uma pesquisa qualitativa do tipo revisão documental comparada, de caráter "
    "descritivo-analítico, adequada ao objetivo de estabelecer a presença ou ausência de um "
    "tópico no currículo prescrito de múltiplos sistemas nacionais: o objeto são documentos "
    "oficiais, e a comparação sistemática entre eles é o método próprio para evidenciar "
    "padrões e exceções. Foram analisados documentos curriculares e de avaliação de sistemas "
    "nacionais dos cinco continentes — de alta, média e baixa renda — e do programa "
    "International Baccalaureate, catalogando-se, para cada sistema, a estrutura do ensino "
    "médio, a presença ou ausência de limites, derivadas e integrais e a fonte oficial "
    "correspondente. Adotou-se uma definição operacional explícita de “Cálculo no currículo” "
    "(aparição explícita de limite, derivada ou integral como conteúdo prescrito), o que torna "
    "a codificação transparente e contestável. Critérios de qualidade: rastreabilidade e "
    "verificação das fontes (cada documento oficial identificado por URL e soma de verificação "
    "SHA-256, com data de captura); triangulação entre currículo prescrito e documentos de "
    "avaliação; e — diferencial do trabalho — reprodutibilidade automatizada, mediante roteiro "
    "que extrai e confere o achado central sobre os documentos, permitindo a terceiros replicar "
    "a verificação em minutos. O mérito metodológico está nessa auditabilidade: o achado central "
    "não depende da autoridade dos autores, mas de evidência documental pública e reexecutável."
)
B4 = (
    "O estudo é uma revisão de fontes curriculares e de avaliação públicas e oficiais; não "
    "envolve seres humanos, dados pessoais ou experimentação, não requerendo apreciação por "
    "Comitê de Ética em Pesquisa. Em atenção à política de transparência do periódico, "
    "declara-se o uso do assistente de inteligência artificial Claude (Opus 4.8, Anthropic) "
    "como ferramenta de apoio em todas as etapas (curadoria e revisão de fontes, redação e "
    "edição, geração do código das figuras), com revisão e verificação humanas integrais e "
    "responsabilidade pública dos autores pelo conteúdo — declaração também registrada no "
    "manuscrito."
)
B5 = (
    "As principais limitações são: (i) documentos curriculares expressam o currículo prescrito "
    "(desiderato), não o currículo implementado em sala, de modo que a análise não captura "
    "práticas efetivas de ensino; (ii) o argumento não é de inferência causal estrita — não se "
    "afirma que a ausência de Cálculo cause o desempenho observado; os indicadores contextuais "
    "(PISA e reprovação em Cálculo I) são apresentados como associações, e a correlação entre "
    "status curricular e PISA é reportada com honestidade, inclusive quando fraca ou nula, com "
    "o Brasil tratado como caso extremo; (iii) a amostra de sistemas, embora ampla e diversa, é "
    "intencional, e não exaustiva."
)
C6 = (
    "A contribuição original e específica é tripla. Primeiro, um mapeamento internacional "
    "inédito e preciso que documenta, sobre fontes oficiais, que o Brasil é o único país da "
    "amostra cujo currículo nacional do ensino médio exclui, simultânea e integralmente, o "
    "Cálculo diferencial e o integral — com o dado contraintuitivo de que vários países de IDH "
    "inferior ao brasileiro (Índia, Vietnã, Paquistão, Nigéria) oferecem Cálculo completo, o "
    "que afasta a explicação por renda e devolve a questão ao terreno curricular e pedagógico. "
    "Segundo, a leitura teórica dessa ausência à luz da transposição didática de Chevallard, "
    "como não-transposição do saber de referência ao saber a ensinar. Terceiro, uma ponte para "
    "a sala de aula: o desenho e a análise a priori, na tradição da engenharia didática, de uma "
    "sequência para introduzir derivada e integral no ensino médio segundo a lógica "
    "Concrete-Pictorial-Abstract, cuja validação empírica se explicita como etapa seguinte da "
    "pesquisa."
)
C7 = (
    "O trabalho avança o conhecimento ao converter uma intuição recorrente da área em achado "
    "sistematizado, verificável e delimitado. Oferece três novidades: (i) um instrumento "
    "reproduzível — manifesto público de fontes oficiais com hash SHA-256 e roteiro de "
    "verificação — que traz ao gênero da pesquisa curricular um padrão de auditabilidade raro; "
    "(ii) um recorte analítico específico (presença ou ausência do Cálculo antes do ensino "
    "superior) que complementa e tensiona os estudos curriculares comparados já publicados no "
    "periódico, ampliando-os do plano da estrutura para o do conteúdo; e (iii) uma proposta "
    "didática fundamentada, que não se limita ao diagnóstico. Em relação a estudos anteriores, "
    "os resultados confirmam e precisam, com evidência documental país a país, o que a "
    "literatura brasileira afirmava de modo qualitativo, e contrariam a explicação simples pelo "
    "nível de desenvolvimento econômico."
)
C8 = (
    "A relevância é direta para a comunidade de Educação Matemática e para o debate curricular "
    "brasileiro em curso (BNCC e Novo Ensino Médio), ao oferecer uma referência comparada firme "
    "sobre o que fazem os demais sistemas nacionais. As implicações são teóricas (a categoria de "
    "não-transposição didática) e práticas (a sequência em engenharia didática, à disposição de "
    "professores e formadores). Interessam especificamente a pesquisadores de currículo e de "
    "didática do Cálculo, a formuladores de políticas curriculares, a formadores de professores "
    "e às licenciaturas e engenharias — estas diretamente afetadas pelas elevadas taxas de "
    "reprovação em Cálculo I. O impacto social e econômico potencial associa-se à formação "
    "matemática de base para as carreiras científicas e tecnológicas e à equidade, uma vez que a "
    "ausência do Cálculo no currículo público tende a ampliar desigualdades de acesso ao ensino "
    "superior de exatas. Espera-se, por fim, contribuir para o fortalecimento e a diversificação "
    "da pesquisa na área, ao acoplar, em um mesmo estudo, análise curricular comparada, "
    "fundamentação didática e reprodutibilidade auditável."
)

# âncora (texto exato) -> resposta, inserida logo após a questão-âncora
ANSWERS = {
    "Cite 5-8 referências fundamentais que delimitam o estado atual do conhecimento na área.": A1,
    "Como essas lacunas justificam a necessidade de sua pesquisa?": A2,
    "Qual é o mérito metodológico (no caso de um manuscrito metodológico/empírico)?": B3,
    "Como as questões éticas envolvidas foram consideradas?": B4,
    "Quais são as principais limitações de seu estudo?": B5,
    "Como seus resultados se diferenciam, ampliam, contrariam ou complementam estudos anteriores?": C7,
    "Qual o impacto esperado do manuscrito proposto para o fortalecimento e/ou diversificação da pesquisa na área de Educação Matemática?": C8,
}

# helpers -----------------------------------------------------------------
def make_par(after_p, text, bold=False, italic=False, align=AL.JUSTIFY):
    new = OxmlElement("w:p")
    after_p._p.addnext(new)
    par = Paragraph(new, after_p._parent)
    par.style = d.styles["Normal"]
    par.alignment = align
    par.paragraph_format.space_after = Pt(6)
    r = par.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    r.bold, r.italic = bold, italic
    # rFonts completo
    rpr = r._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = rpr.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, rf)
    for a in ("w:ascii", "w:hAnsi", "w:cs"):
        rf.set(qn(a), "Times New Roman")
    return par


# 1) capturar referências das questões-âncora ANTES de mexer no doc
targets = {}
para23 = None
for p in d.paragraphs:
    t = p.text.strip()
    if t in ANSWERS and t not in targets:
        targets[t] = p
    if t.startswith("Sua contribuição é"):
        para23 = p

# item 6 (contribuição): inserir após a linha de checkboxes (para23)
# 2) remover blocos EN + ES (do "Mandatory Structure" em diante)
remove = False
for p in list(d.paragraphs):
    if p.text.strip().startswith("BOLEMA – Mandatory Structure"):
        remove = True
    if remove:
        p._element.getparent().remove(p._element)

# 3) marcar tipo de contribuição e inserir C6 após os checkboxes
if para23 is not None:
    for r in list(para23.runs):
        r.text = ""
    run = para23.runs[0] if para23.runs else para23.add_run("")
    run.text = ("Sua contribuição é (assinale uma ou mais opções): ( X ) Teórica    "
                "(  ) Metodológica    ( X ) Empírica    ( X ) Prática/Aplicada    "
                "(  ) Combinação (especifique).")
    run.font.name = "Times New Roman"; run.font.size = Pt(12)
    make_par(para23, C6)

# 4) inserir as demais respostas após suas âncoras
for anchor, ans in ANSWERS.items():
    make_par(targets[anchor], ans)

# 5) cabeçalho de identificação logo após o título
title_p = d.paragraphs[0]
make_par(title_p,
         "Manuscrito: “O cálculo ausente: duzentos anos de currículo, uma comparação "
         "internacional e um roteiro pedagógico para o ensino médio” — Tipo: Artigo (original). "
         "Idioma: Português.", italic=True)

d.core_properties.author = "Leonardo Chalhoub"
d.core_properties.title = "Cover Letter — O cálculo ausente"
d.save(OUT)

# relatório de contagem por seção
import re
secwords = {"A": 0, "B": 0, "C": 0}
cur = None
for p in Document(OUT).paragraphs:
    t = p.text.strip()
    if t.startswith("Seção A"): cur = "A"
    elif t.startswith("Seção B"): cur = "B"
    elif t.startswith("Seção C"): cur = "C"
    if cur:
        secwords[cur] += len(t.split())
print("OK ->", OUT)
print("palavras/seção (inclui perguntas):", secwords, "| total:", sum(secwords.values()))
